"""Tests de los loaders de borradores."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from blog_auditor.exceptions import LoaderError
from blog_auditor.loaders import gdoc_loader, load_draft_text, pdf_loader
from blog_auditor.loaders.gdoc_loader import is_google_doc_url, load_google_doc

# ── Dispatch ────────────────────────────────────────────────────


def test_empty_source_raises() -> None:
    with pytest.raises(LoaderError, match="vacía"):
        load_draft_text("   ")


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(LoaderError, match="No existe"):
        load_draft_text(str(tmp_path / "no-existe.md"))


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    file = tmp_path / "borrador.pptx"
    file.write_text("contenido", encoding="utf-8")

    with pytest.raises(LoaderError, match="Formato no soportado"):
        load_draft_text(str(file))


def test_non_gdoc_url_raises() -> None:
    with pytest.raises(LoaderError, match="Google Docs"):
        load_draft_text("https://teilurtalent.com/insights/algun-articulo")


# ── Texto plano / Markdown ──────────────────────────────────────


def test_loads_markdown_file(tmp_path: Path) -> None:
    file = tmp_path / "borrador.md"
    file.write_text("# Título\n\nContenido del artículo.", encoding="utf-8")

    assert load_draft_text(str(file)) == "# Título\n\nContenido del artículo."


def test_loads_utf8_with_bom(tmp_path: Path) -> None:
    file = tmp_path / "borrador.txt"
    file.write_bytes("Artículo con BOM".encode("utf-8-sig"))

    assert load_draft_text(str(file)) == "Artículo con BOM"


def test_falls_back_to_cp1252(tmp_path: Path) -> None:
    file = tmp_path / "borrador.txt"
    file.write_bytes("Café con leche".encode("cp1252"))

    assert load_draft_text(str(file)) == "Café con leche"


# ── Word (.docx) ────────────────────────────────────────────────


def test_loads_docx_preserving_headings(tmp_path: Path) -> None:
    path = tmp_path / "borrador.docx"
    document = Document()
    document.add_heading("Toptal Pricing Explained", level=1)
    document.add_paragraph("Toptal does not publish its rates.")
    document.add_heading("How the model works", level=2)
    document.add_paragraph("They charge a hidden margin.")
    document.save(str(path))

    text = load_draft_text(str(path))

    assert "# Toptal Pricing Explained" in text
    assert "## How the model works" in text
    assert "Toptal does not publish its rates." in text


def test_invalid_docx_raises(tmp_path: Path) -> None:
    path = tmp_path / "falso.docx"
    path.write_text("esto no es un docx", encoding="utf-8")

    with pytest.raises(LoaderError, match="no es un .docx válido"):
        load_draft_text(str(path))


# ── PDF ─────────────────────────────────────────────────────────


class FakePdfPage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


class FakePdfReader:
    """Doble de PdfReader para probar la lógica del loader sin crear PDFs."""

    pages_by_path: dict[str, list[FakePdfPage]] = {}
    encrypted = False

    def __init__(self, path: str) -> None:
        self.pages = self.pages_by_path.get(path, [])
        self.is_encrypted = self.encrypted

    def decrypt(self, password: str) -> None:
        raise NotImplementedError("algoritmo no soportado")


@pytest.fixture()
def fake_pdf_reader(monkeypatch: pytest.MonkeyPatch) -> type[FakePdfReader]:
    FakePdfReader.pages_by_path = {}
    FakePdfReader.encrypted = False
    monkeypatch.setattr(pdf_loader, "PdfReader", FakePdfReader)
    return FakePdfReader


def test_loads_pdf_normalizing_whitespace(
    tmp_path: Path, fake_pdf_reader: type[FakePdfReader]
) -> None:
    path = tmp_path / "borrador.pdf"
    path.write_bytes(b"%PDF-fake")
    fake_pdf_reader.pages_by_path[str(path)] = [
        FakePdfPage("How  to  Hire  DevOps  Engineers  Without\nSlowing\n \nDown\n \n"),
        FakePdfPage("Second  page  content."),
    ]

    text = load_draft_text(str(path))

    assert "How to Hire DevOps Engineers Without" in text
    assert "  " not in text
    assert "Second page content." in text


def test_pdf_without_text_raises(tmp_path: Path, fake_pdf_reader: type[FakePdfReader]) -> None:
    path = tmp_path / "escaneo.pdf"
    path.write_bytes(b"%PDF-fake")
    fake_pdf_reader.pages_by_path[str(path)] = [FakePdfPage("   \n ")]

    with pytest.raises(LoaderError, match="texto extraíble"):
        load_draft_text(str(path))


def test_encrypted_pdf_raises(tmp_path: Path, fake_pdf_reader: type[FakePdfReader]) -> None:
    path = tmp_path / "protegido.pdf"
    path.write_bytes(b"%PDF-fake")
    fake_pdf_reader.encrypted = True

    with pytest.raises(LoaderError, match="contraseña"):
        load_draft_text(str(path))


def test_invalid_pdf_raises(tmp_path: Path) -> None:
    path = tmp_path / "falso.pdf"
    path.write_text("esto no es un pdf", encoding="utf-8")

    with pytest.raises(LoaderError, match="No pude abrir el PDF"):
        load_draft_text(str(path))


# ── Google Docs ─────────────────────────────────────────────────


class FakeResponse:
    def __init__(self, status_code: int, text: str) -> None:
        self.status_code = status_code
        self.text = text


GDOC_URL = "https://docs.google.com/document/d/abc123XYZ_-/edit?usp=sharing"


def test_is_google_doc_url() -> None:
    assert is_google_doc_url(GDOC_URL)
    assert not is_google_doc_url("https://teilurtalent.com")


def test_loads_public_google_doc(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_get(url: str, **kwargs: object) -> FakeResponse:
        assert "abc123XYZ_-" in url
        assert "export?format=txt" in url
        return FakeResponse(200, "Contenido del borrador desde Google Docs.")

    monkeypatch.setattr(gdoc_loader.httpx, "get", fake_get)

    assert load_google_doc(GDOC_URL) == "Contenido del borrador desde Google Docs."


def test_private_doc_raises_clear_error(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(gdoc_loader.httpx, "get", lambda url, **kw: FakeResponse(403, ""))

    with pytest.raises(LoaderError, match="compartido"):
        load_google_doc(GDOC_URL)


def test_login_page_raises_clear_error(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        gdoc_loader.httpx,
        "get",
        lambda url, **kw: FakeResponse(200, "<html><body>Sign in</body></html>"),
    )

    with pytest.raises(LoaderError, match="login"):
        load_google_doc(GDOC_URL)


def test_empty_doc_raises(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(gdoc_loader.httpx, "get", lambda url, **kw: FakeResponse(200, "   "))

    with pytest.raises(LoaderError, match="vacío"):
        load_google_doc(GDOC_URL)
